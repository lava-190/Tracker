import tkinter as tk
from tkinter import messagebox, colorchooser, filedialog, simpledialog
import datetime
import pandas as pd
import os
import json
from docx import Document
from fpdf import FPDF
import shutil
import arabic_reshaper
from bidi.algorithm import get_display
from PIL import Image, ImageTk
import random
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

def reshape_arabic_text(text):
    reshaped_text = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped_text)
    return bidi_text

# إعداد المسارات والملفات الأساسية
LISTS_FILE = "lists.json"
MAIN_FOLDER = "Lists"

# ---------------------------
# دوال المساعدة لتحميل وحفظ البيانات
# ---------------------------
def load_lists():
    if os.path.exists(LISTS_FILE):
        with open(LISTS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_lists(lists_data):
    with open(LISTS_FILE, "w", encoding="utf-8") as f:
        json.dump(lists_data, f, ensure_ascii=False, indent=4)

def backup_data():
    backup_folder = "Backup"
    if not os.path.exists(backup_folder):
        os.makedirs(backup_folder)
    if os.path.exists(LISTS_FILE):
        shutil.copy(LISTS_FILE, os.path.join(backup_folder, LISTS_FILE))
        messagebox.showinfo("Backup", "تم عمل النسخة الاحتياطية بنجاح!")
    else:
        messagebox.showerror("Backup", "لم يتم العثور على ملف البيانات للنسخ الاحتياطي.")

# ---------------------------
# فئة ToolTip لإظهار التلميحات عند مرور الماوس
# ---------------------------
class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tipwindow = None
        widget.bind("<Enter>", self.show_tip)
        widget.bind("<Leave>", self.hide_tip)

    def show_tip(self, event=None):
        if self.tipwindow or not self.text:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        self.tipwindow = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry("+%d+%d" % (x, y))
        label = tk.Label(tw, text=self.text, background="#ffffe0", relief="solid", borderwidth=1, font=("Arial", 10))
        label.pack(ipadx=1)

    def hide_tip(self, event=None):
        if self.tipwindow:
            self.tipwindow.destroy()
        self.tipwindow = None

# ---------------------------
# نافذة عرض التقدم (تقارير)
# ---------------------------
class ProgressWindow(tk.Toplevel):
    def __init__(self, master, list_name, tasks):
        super().__init__(master)
        self.list_name = list_name
        self.tasks = tasks
        self.list_folder = os.path.join(MAIN_FOLDER, list_name)
        self.excel_file = os.path.join(self.list_folder, "daily_progress.xlsx")
        self.title(f"عرض التقدم - {list_name}")
        self.geometry("700x600")
        self.create_widgets()
    
    def create_widgets(self):
        tk.Label(self, text=f"عرض التقدم للقائمة: {self.list_name}", font=("Arial", 16)).pack(pady=10)
        
        tk.Button(self, text="عرض التقدم اليومي", command=self.show_daily_progress).pack(pady=5)
        rep_frame = tk.Frame(self)
        rep_frame.pack(pady=10)
        tk.Button(rep_frame, text="توليد التقرير الأسبوعي", command=lambda: self.generate_report(period="weekly")).pack(side="left", padx=5)
        tk.Button(rep_frame, text="توليد التقرير الشهري", command=lambda: self.generate_report(period="monthly")).pack(side="left", padx=5)
        tk.Button(rep_frame, text="توليد تقرير PDF", command=self.generate_pdf_report).pack(side="left", padx=5)
        tk.Button(self, text="عرض التقرير التفاعلي", command=self.interactive_report).pack(pady=5)
    
    def show_daily_progress(self):
        if not os.path.exists(self.excel_file):
            messagebox.showerror("خطأ", "لا توجد بيانات يومية متوفرة.")
            return
        try:
            df = pd.read_excel(self.excel_file)
            today_str = datetime.date.today().strftime("%Y-%m-%d")
            df_today = df[df['التاريخ'] == today_str]
            if df_today.empty:
                messagebox.showinfo("التقدم اليومي", "لا توجد بيانات ليومنا هذا.")
                return
            record = df_today.iloc[-1]
            progress_text = f"التاريخ: {record['التاريخ']}\n\n"
            for task_obj in self.tasks:
                task = task_obj["task"]
                status = record.get(task, "✖️")
                comment = record.get(f"{task}_تعليق", "")
                progress_text += f"المهمة: {task}\nالحالة: {status}\nالتعليق: {comment}\n\n"
            progress_window = tk.Toplevel(self)
            progress_window.title("التقدم اليومي")
            text = tk.Text(progress_window, wrap="none", font=("Arial", 10))
            text.pack(fill="both", expand=True)
            text.insert(tk.END, progress_text)
        except Exception as e:
            messagebox.showerror("خطأ", f"حدث خطأ أثناء قراءة البيانات:\n{e}")
    
    def generate_report(self, period="weekly"):
        if not os.path.exists(self.excel_file):
            messagebox.showerror("خطأ", "لا توجد بيانات.")
            return
        df = pd.read_excel(self.excel_file)
        df['التاريخ'] = pd.to_datetime(df['التاريخ'])
        today = pd.Timestamp.today().normalize()
        if period == "weekly":
            start_period = today - pd.Timedelta(days=today.weekday())
            folder_name = "Weekly_Reports"
        else:
            start_period = today.replace(day=1)
            folder_name = "Monthly_Reports"
        period_data = df[df['التاريخ'] >= start_period]
        if period_data.empty:
            messagebox.showinfo("تقرير", f"لا توجد بيانات للفترة المحددة ({period}).")
            return
        summary = {}
        for task_obj in self.tasks:
            task = task_obj["task"]
            count = period_data[task].apply(lambda x: 1 if x == "✔" else 0).sum()
            comments_series = period_data[f"{task}_تعليق"].dropna().astype(str)
            comments = "; ".join([c for c in comments_series if c.strip() != ""])
            summary[task] = {"count": count, "comments": comments}
        document = Document()
        document.add_heading(f"تقرير {period} - {self.list_name}", 0)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "المهمة"
        hdr_cells[1].text = "عدد مرات القيام"
        hdr_cells[2].text = "التعليقات"
        for task_obj in self.tasks:
            task = task_obj["task"]
            row_cells = table.add_row().cells
            row_cells[0].text = task
            row_cells[1].text = str(summary[task]["count"])
            row_cells[2].text = summary[task]["comments"]
        report_date = today.strftime("%Y-%m-%d")
        report_folder = os.path.join(self.list_folder, folder_name)
        if not os.path.exists(report_folder):
            os.makedirs(report_folder)
        word_report_file = os.path.join(report_folder, f"{period}_report_{report_date}.docx")
        document.save(word_report_file)
        messagebox.showinfo("تقرير", f"تم حفظ تقرير {period} في:\n{word_report_file}")

    def generate_pdf_report(self):
        if not os.path.exists(self.excel_file):
            messagebox.showerror("خطأ", "لا توجد بيانات.")
            return
        df = pd.read_excel(self.excel_file)
        df['التاريخ'] = pd.to_datetime(df['التاريخ'])
        today = pd.Timestamp.today().normalize()
        start_date = today - pd.Timedelta(days=7)
        period_data = df[df['التاريخ'] >= start_date]
        if period_data.empty:
            messagebox.showinfo("تقرير", "لا توجد بيانات للفترة المحددة.")
            return
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', '', 14)
         
        title = reshape_arabic_text(f"تقرير PDF للمهام - {self.list_name}")
        pdf.cell(200, 10, txt=title, ln=True, align="C")
        pdf.ln(10)
         
        for task_obj in self.tasks:
            task_original = task_obj["task"]
            count = period_data[task_original].apply(lambda x: 1 if x == "✔" else 0).sum()
            comments_series = period_data[f"{task_original}_تعليق"].dropna().astype(str)
            comments = "; ".join([c for c in comments_series if c.strip() != ""])
    
            pdf.set_font('DejaVu', '', 12)
            pdf.cell(0, 10, txt=reshape_arabic_text(f"المهمة: {task_original}"), ln=True)
            pdf.cell(0, 10, txt=reshape_arabic_text(f"عدد المرات: {count}"), ln=True)
            pdf.cell(0, 10, txt=reshape_arabic_text(f"التعليقات: {comments}"), ln=True)
    
            pdf.ln(10)
 
        pdf_folder = os.path.join(self.list_folder, "PDF_Reports")
        if not os.path.exists(pdf_folder):
            os.makedirs(pdf_folder)
        pdf_file = os.path.join(pdf_folder, f"pdf_report_{today.strftime('%Y-%m-%d')}.pdf")
        pdf.output(pdf_file)
        messagebox.showinfo("تقرير", f"تم حفظ تقرير PDF في:\n{pdf_file}")

    def interactive_report(self):
        import matplotlib.font_manager as fm
        import matplotlib as mpl
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

        # إذا لم تكن قد استدعيت الخط خارج الدالة، يمكنك استدعاؤه هنا أيضاً
        font_path = "DejaVuSans.ttf"
        font_prop = fm.FontProperties(fname=font_path)
        mpl.rcParams['font.family'] = font_prop.get_name()
        mpl.rcParams['axes.unicode_minus'] = False

        if not os.path.exists(self.excel_file):
            messagebox.showerror("خطأ", "لا توجد بيانات.")
            return
    
        df = pd.read_excel(self.excel_file)
        df['التاريخ'] = pd.to_datetime(df['التاريخ'])
        today = pd.Timestamp.today().normalize()
        start_date = today - pd.Timedelta(days=7)
        period_data = df[df['التاريخ'] >= start_date]
        if period_data.empty:
            messagebox.showinfo("تقرير", "لا توجد بيانات للفترة المحددة.")
            return

        summary = {}
        for task_obj in self.tasks:
            task = task_obj["task"]
            count = period_data[task].apply(lambda x: 1 if x == "✔" else 0).sum()
            summary[task] = count

        report_win = tk.Toplevel(self)
        report_win.title("التقرير التفاعلي")

        fig, ax = plt.subplots(figsize=(6,4))

        # إعادة تشكيل أسماء المهام قبل عرضها
        tasks = [reshape_arabic_text(t) for t in summary.keys()]
        counts = list(summary.values())

        ax.bar(tasks, counts, color="skyblue")
        ax.set_title(reshape_arabic_text("عدد مرات إنجاز المهام"))
        ax.set_ylabel(reshape_arabic_text("العدد"))
        ax.tick_params(axis='x', rotation=45)

        fig.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=report_win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)


# ---------------------------
# التطبيق الرئيسي لإدارة القوائم
# ---------------------------
class TaskManagerApp(tk.Tk):
    def show_help(self):
        help_text = (
        
            "مرحباً بك في مدير القوائم!\n\n"
            "هدف التطبيق:\n"
            "تم تصميم مدير القوائم لمساعدتك في إنشاء وإدارة قوائم المهام بسهولة، مع تتبع الأداء اليومي وتوليد التقارير لمتابعة تقدمك. "
            "يوفر التطبيق تجربة مرنة من خلال القائمة الجانبية والنقر بزر الماوس الأيمن للوصول السريع إلى مختلف الوظائف.\n\n"
    
            "كيفية الاستخدام:\n"
            "1. القائمة الجانبية:\n"
            "   - استخدم الأزرار الجانبية لإنشاء، فتح، تعديل، أو حذف القوائم.\n"
            "   - يمكنك ترتيب القوائم باستخدام الأسهم، وتغيير ألوانها لتسهيل التمييز بينها.\n"
            "   - يمكنك عمل نسخة احتياطية من بياناتك بنقرة واحدة.\n"
            "   - يمكن تغيير حجم الخط والخلفية من خلال إعدادات التخصيص.\n\n"
    
            "2. القائمة السياقية (النقر بزر الماوس الأيمن):\n"
            "   - انقر بزر الماوس الأيمن على أي قائمة لعرض الخيارات التالية:\n"
            "     • فتح القائمة لعرض المهام وسجل الأداء.\n"
            "     • تعديل اسم القائمة أو تغيير لونها.\n"
            "     • حذف القائمة أو تصديرها.\n"
            "     • عرض التقدم والتقارير، وإعادة ترتيب القوائم بسهولة.\n\n"
    
            "3. تتبع الأداء وتوليد التقارير:\n"
            "   - عند فتح أي قائمة، يمكنك تسجيل الإنجازات اليومية عبر مربعات الاختيار.\n"
            "   - يمكنك إضافة ملاحظات حول كل مهمة.\n"
            "   - زر 'عرض التقدم' يتيح لك توليد تقارير أسبوعية، شهرية، وPDF لمتابعة الأداء.\n\n"
    
            "4. تخصيص الواجهة:\n"
            "   - يمكنك تعديل حجم الخط واختيار ألوان القوائم من خلال قائمة التخصيص.\n"
            "   - يدعم التطبيق الوضع الليلي وتغيير الخلفية حسب رغبتك.\n\n"
    
            "استمتع باستخدام مدير القوائم مع تجربة أكثر مرونة وسهولة في التنقل!"
        )
        messagebox.showinfo("دليل المستخدم", help_text)


    def __init__(self):
        super().__init__()
        self.title("مدير القوائم")
        self.geometry("1200x550")
        self.dark_mode = False
        self.lists_data = load_lists()
        self.bg_type = None  
        self.bg_value = None
        self.font_size = 12  # إعداد افتراضي لحجم الخط
        self.background_label = None
        # متغيرات جديدة لتخزين ألوان وترتيب القوائم
        self.lists_colors = {}
        self.lists_order = list(self.lists_data.keys())
        self.load_config()
        self.create_context_menu()
        self.create_widgets()
        self.bind("<Configure>", self.on_resize)
       
    def create_context_menu(self):
        self.context_menu = tk.Menu(self, tearoff=0)
        # الخيار الأول: فتح القائمة عبر نافذة اختيار القائمة
        self.context_menu.add_command(label="فتح القائمة", command=self.open_list)
        self.context_menu.add_command(label="تعديل اسم القائمة", command=self.rename_list)
        self.context_menu.add_command(label="تغيير لون القائمة", command=self.change_list_color)
        self.context_menu.add_command(label="تعديل القائمة", command=self.edit_list)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="إنشاء قائمة جديدة", command=self.create_new_list)
        self.context_menu.add_command(label="حذف القائمة", command=self.delete_list)
        self.context_menu.add_command(label="تصدير القائمة", command=self.export_list)
        self.context_menu.add_command(label="عرض التقدم", command=self.open_progress)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="↑", command=self.move_list_up)
        self.context_menu.add_command(label="↓", command=self.move_list_down)



    def clear_widgets(self):
        for widget in self.winfo_children():
            widget.destroy()     

    def load_config(self):
        config_file = "config.json"
        if os.path.exists(config_file):
            with open(config_file, "r", encoding="utf-8") as f:
                config = json.load(f)
            self.bg_type = config.get("bg_type")
            self.bg_value = config.get("bg_value")
            self.font_size = config.get("font_size", 12)
            self.lists_colors = config.get("lists_colors", {})
            self.lists_order = config.get("lists_order", list(self.lists_data.keys()))
        else:
            self.lists_colors = {}
            self.lists_order = list(self.lists_data.keys())

    def save_config(self):
        config_file = "config.json"
        config = {
            "bg_type": self.bg_type,
            "bg_value": self.bg_value,
            "font_size": self.font_size,
            "lists_colors": self.lists_colors,
            "lists_order": self.lists_order
        }
        with open(config_file, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=4)

    def rename_list(self):
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة لتعديل اسمها.")
            return
        old_name = self.lists_listbox.get(selection[0])
        new_name = simpledialog.askstring("تعديل اسم القائمة", "أدخل الاسم الجديد للقائمة:", initialvalue=old_name)
        if new_name and new_name.strip() and new_name != old_name:
            if new_name in self.lists_data:
                messagebox.showerror("خطأ", "هذه القائمة موجودة بالفعل.")
                return
            self.lists_data[new_name] = self.lists_data.pop(old_name)
            save_lists(self.lists_data)
            old_folder = os.path.join(MAIN_FOLDER, old_name)
            new_folder = os.path.join(MAIN_FOLDER, new_name)
            if os.path.exists(old_folder):
                os.rename(old_folder, new_folder)
            if old_name in self.lists_order:
                index = self.lists_order.index(old_name)
                self.lists_order[index] = new_name
            if old_name in self.lists_colors:
                self.lists_colors[new_name] = self.lists_colors.pop(old_name)
            messagebox.showinfo("نجاح", "تم تعديل اسم القائمة بنجاح!")
            self.save_config()
            self.refresh_lists()

    def create_widgets(self):
        top_bar = tk.Frame(self)
        top_bar.pack(side="top", fill="x")    
        self.toggle_settings_btn = tk.Button(top_bar, text="≡", font=("Arial", self.font_size+4), command=self.toggle_settings)
        self.toggle_settings_btn.pack(side="left", padx=5, pady=5)
    
        self.side_menu = tk.Frame(self, bd=2, relief="raised")
        self.side_menu.pack(side="left", fill="y")
        self.main_frame = tk.Frame(self)
        self.main_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)
    
        tk.Label(self.side_menu, text="القوائم", font=("Arial", self.font_size, "bold")).pack(pady=5)
    
        btn_new = tk.Button(self.side_menu, text="إنشاء قائمة جديدة", command=self.create_new_list, width=20)
        btn_new.pack(pady=2)
        ToolTip(btn_new, "أنشئ قائمة جديدة لإضافة مهامك")

        btn_open = tk.Button(self.side_menu, text="فتح القائمة", command=lambda: self.select_list_and_execute(self.open_list_by_name), width=20)
        btn_open.pack(pady=2)
        ToolTip(btn_open, "اختر قائمة لفتحها")

        btn_edit = tk.Button(self.side_menu, text="تعديل القائمة", command=lambda: self.select_list_and_execute(self.edit_list_by_name), width=20)        
        btn_edit.pack(pady=2)
        ToolTip(btn_edit, "اختر قائمة لتعديلها")
    
        btn_delete = tk.Button(self.side_menu, text="حذف القائمة", command=lambda: self.select_list_and_execute(self.delete_list_by_name), width=20)        
        btn_delete.pack(pady=2)
        ToolTip(btn_delete, "اختر قائمة لحذفها")
    
        btn_export = tk.Button(self.side_menu, text="تصدير القائمة", command=lambda: self.select_list_and_execute(self.export_list_by_name), width=20)
        btn_export.pack(pady=2)
        ToolTip(btn_export, "اختر قائمة للتصدير")
       
        btn_color = tk.Button(self.side_menu, text="تلوين القوائم", command=lambda: self.select_list_and_execute(self.change_list_color_by_name) , width=20)
        btn_color.pack(pady=2)
        ToolTip(btn_color, "اختر قائمة لتغيير لونها")


        btn_reorder = tk.Button(self.side_menu, text="ترتيب القوائم", command=self.reorder_lists_window, width=20)
        btn_reorder.pack(pady=2)
        ToolTip(btn_reorder, "إعادة ترتيب القوائم")




        btn_backup = tk.Button(self.side_menu, text="نسخ احتياطي", command=backup_data, width=20)
        btn_backup.pack(pady=2)
        ToolTip(btn_backup, "اعمل نسخة احتياطية من بياناتك")
    
        btn_progress = tk.Button(self.side_menu, text="عرض التقدم", command=lambda: self.select_list_and_execute(self.open_progress_by_name), width=20)
        btn_progress.pack(pady=2)
        ToolTip(btn_progress, "اختر قائمة لعرض التقدم")    
        tk.Label(self.side_menu, text="الخلفية", font=("Arial", self.font_size, "bold")).pack(pady=10)
    
        btn_toggle = tk.Button(self.side_menu, text="تبديل الوضع الليلي", command=self.toggle_dark_mode, width=20)
        btn_toggle.pack(pady=2)
        ToolTip(btn_toggle, "تبديل بين الوضع الليلي والعادي")
    
        btn_bg = tk.Button(self.side_menu, text="تغيير الخلفية (ألوان)", command=self.change_background, width=20)
        btn_bg.pack(pady=2)
        ToolTip(btn_bg, "اختر لون خلفية للتطبيق")
    
        btn_remove_bg = tk.Button(self.side_menu, text="إزالة الخلفية", command=self.remove_background, width=20)
        btn_remove_bg.pack(pady=2)
        ToolTip(btn_remove_bg, "إزالة الخلفية المخصصة")
    
        btn_customize = tk.Button(self.side_menu, text="تخصيص الواجهة", command=self.customize_ui, width=20)
        btn_customize.pack(pady=2)
        ToolTip(btn_customize, "تغيير حجم الخط والثيمات")
    
        main_frame = tk.Frame(self)
        main_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)
    
        top_frame = tk.Frame(main_frame)
        top_frame.pack(pady=10, fill="x")
        tk.Label(top_frame, text="بحث:", font=("Arial", self.font_size)).pack(side="left", padx=5)
    
        self.search_entry = tk.Entry(top_frame, font=("Arial", self.font_size))
        self.search_entry.pack(side="left", padx=5)
        self.search_entry.bind("<KeyRelease>", lambda event: self.refresh_lists())
    
        btn_help = tk.Button(top_frame, text="?", font=("Arial", self.font_size, "bold"), command=self.show_help)
        btn_help.pack(side="right", padx=5)
        ToolTip(btn_help, "دليل الاستخدام")
    
        tk.Label(top_frame, text="صلِّ على النبي", font=("Arial", self.font_size), fg="blue").pack(side="right", padx=5)
    
        tk.Label(main_frame, text="القوائم المتاحة", font=("Arial", self.font_size+4)).pack(pady=5)
        default_bg = self.bg_value if self.bg_type=="color" and self.bg_value else ("black" if self.dark_mode else "white")
        default_fg = "white" if self.dark_mode else "black"
        self.lists_listbox = tk.Listbox(
            main_frame,
            height=8,
            font=("Arial", self.font_size),
            activestyle="none",
            exportselection=False,
            selectbackground=default_bg,
            selectforeground=default_fg,
            highlightthickness=0,
            bd=0
        )
        self.lists_listbox.bind("<<ListboxSelect>>", lambda e: self.lists_listbox.selection_clear(0, tk.END))
        self.lists_listbox.pack(fill="both", expand=True, padx=20, pady=10)
        self.lists_listbox.bind("<Button-3>", self.show_context_menu)
        self.lists_listbox.bind("<Button-2>", self.show_context_menu)
        
        
        self.refresh_lists()
    
        tk.Label(self, text="created by: meedoasadel@gmail.com", font=("Arial", 20, "bold"), fg="blue").pack(side="bottom", pady=15)
        self.apply_theme()
        
    def toggle_settings(self):
        if self.side_menu.winfo_ismapped():
            self.side_menu.pack_forget()
        else:
            self.side_menu.pack(side="left", fill="y", before=self.main_frame)

    def show_context_menu(self, event):
        try:
            index = self.lists_listbox.nearest(event.y)
            self.lists_listbox.selection_clear(0, tk.END)
            self.lists_listbox.selection_set(index)
            self.lists_listbox.activate(index)
            self.context_menu.post(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()

    def change_list_color(self):
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة لتغيير لونها.")
            return
        index = selection[0]
        list_name = self.lists_listbox.get(index)
        color = colorchooser.askcolor(title="اختر لون القائمة")[1]
        if color:
            self.lists_colors[list_name] = color
            self.save_config()
            self.refresh_lists()


    def apply_theme(self):
        if self.bg_type == "color" and self.bg_value:
            self.configure(bg=self.bg_value)
            self.lists_listbox.configure(bg=self.bg_value, fg="white" if self.dark_mode else "black")
            self.search_entry.configure(bg=self.bg_value, fg="white" if self.dark_mode else "black")
            for widget in self.winfo_children():
                if isinstance(widget, tk.Frame):
                    widget.configure(bg=self.bg_value)
        else:
            bg_color = "black" if self.dark_mode else "white"
            self.configure(bg=bg_color)
            self.lists_listbox.configure(bg=bg_color, fg="white" if self.dark_mode else "black")
            self.search_entry.configure(bg=bg_color, fg="white" if self.dark_mode else "black")
            for widget in self.winfo_children():
                if isinstance(widget, tk.Frame):
                    widget.configure(bg=bg_color)

    def change_background(self):
        color = colorchooser.askcolor(title="اختر لون الخلفية")
        if color[1]:
            self.bg_type = "color"
            self.bg_value = color[1]
            self.apply_theme()
            self.save_config()

    def remove_background(self):
        self.bg_type = None
        self.bg_value = None
        if self.background_label:
            self.background_label.destroy()
            self.background_label = None
        self.apply_theme()
        self.save_config()

    def on_resize(self, event):
        pass

    def toggle_dark_mode(self):
        self.dark_mode = not self.dark_mode
        self.apply_theme()

    def customize_ui(self):
        # نافذة لتخصيص حجم الخط والثيمات
        cust_win = tk.Toplevel(self)
        cust_win.title("تخصيص الواجهة")
        cust_win.geometry("300x200")
        tk.Label(cust_win, text="حجم الخط:", font=("Arial", 12)).pack(pady=10)
        font_entry = tk.Entry(cust_win, font=("Arial", 12))
        font_entry.insert(0, str(self.font_size))
        font_entry.pack(pady=5)

    def reorder_lists_window(self):
        win = tk.Toplevel(self)
        win.title("ترتيب القوائم")
        win.geometry("300x400")
        # إنشاء Listbox يعرض الترتيب الحالي للقوائم
        lb = tk.Listbox(win, font=("Arial", self.font_size))
        lb.pack(fill="both", expand=True, padx=10, pady=10)
        for item in self.lists_order:
            lb.insert(tk.END, item)
        # إطار يحتوي على أزرار التحريك
        btn_frame = tk.Frame(win)
        btn_frame.pack(pady=10)
        def move_up():
            sel = lb.curselection()
            if not sel:
                return
            index = sel[0]
            if index > 0:
                self.lists_order[index], self.lists_order[index-1] = self.lists_order[index-1], self.lists_order[index]
                refresh_order()
        def move_down():
            sel = lb.curselection()
            if not sel:
                return
            index = sel[0]
            if index < lb.size()-1:
                self.lists_order[index], self.lists_order[index+1] = self.lists_order[index+1], self.lists_order[index]
                refresh_order()
        def refresh_order():
            lb.delete(0, tk.END)
            for item in self.lists_order:
                lb.insert(tk.END, item)
            lb.selection_clear(0, tk.END)
            # اختيار العنصر الأول افتراضيًا
            if lb.size() > 0:
                lb.selection_set(0)
        btn_up = tk.Button(btn_frame, text="↑", command=move_up)
        btn_up.pack(side=tk.LEFT, padx=5)
        btn_down = tk.Button(btn_frame, text="↓", command=move_down)
        btn_down.pack(side=tk.LEFT, padx=5)
        # زر حفظ وإغلاق النافذة
        def save_and_close():
            self.save_config()
            self.refresh_lists()
            win.destroy()
        btn_save = tk.Button(win, text="حفظ", command=save_and_close)
        btn_save.pack(pady=5)


        def apply_customization():
            try:
                new_size = int(font_entry.get())
                # فرض الحدود بين 10 و50
                new_size = max(10, min(new_size, 50))
                self.font_size = new_size
                self.save_config()
                # تحديث الواجهة: حذف جميع الواجهات وإعادة إنشائها
                self.clear_widgets()
                self.create_widgets()
                cust_win.destroy()
            except:
                messagebox.showerror("خطأ", "يرجى إدخال رقم صحيح.")
        tk.Button(win, text="تطبيق", command=apply_customization).pack(pady=10)
    
    def move_list_up(self):
        # التأكد من اختيار عنصر من القائمة
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة.")
            return
        index = selection[0]
        # إذا لم يكن العنصر الأول، نقوم بتبديله مع العنصر الذي قبله
        if index > 0:
            self.lists_order[index], self.lists_order[index-1] = self.lists_order[index-1], self.lists_order[index]
            self.save_config()
            self.refresh_lists()
            self.lists_listbox.selection_clear(0, tk.END)
            self.lists_listbox.selection_set(index-1)
    
    def move_list_down(self):
        # التأكد من اختيار عنصر من القائمة
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة.")
            return
        index = selection[0]
        # إذا لم يكن العنصر الأخير، نقوم بتبديله مع العنصر الذي يليه
        if index < len(self.lists_order) - 1:
            self.lists_order[index], self.lists_order[index+1] = self.lists_order[index+1], self.lists_order[index]
            self.save_config()
            self.refresh_lists()
            self.lists_listbox.selection_clear(0, tk.END)
            self.lists_listbox.selection_set(index+1)


    def select_list_and_execute(self, func):
        # نافذة اختيار قائمة لتطبيق العملية المطلوبة (func)
        win = tk.Toplevel(self)
        win.title("اختر القائمة")
        win.geometry("350x350")
        lb = tk.Listbox(win, font=("Arial", self.font_size), selectmode=tk.SINGLE)
        lb.pack(fill="both", expand=True, padx=10, pady=10)
        for name in self.lists_data.keys():
            lb.insert(tk.END, name)
        lb.focus_set()  
        def on_execute():
            selection = lb.curselection()
            if not selection:
                messagebox.showerror("خطأ", "يرجى اختيار قائمة.")
            else:
                list_name = lb.get(selection[0])
                win.destroy()
                func(list_name)
        btn = tk.Button(win, text="تنفيذ", command=on_execute)
        btn.pack(pady=10)


    def open_list_by_name(self, list_name):
        tasks = self.lists_data[list_name]
        DailyTrackerApp(self, list_name, tasks)

    def edit_list_by_name(self, list_name):
        tasks = self.lists_data[list_name]
        CreateListWindow(self, list_name, tasks)

    def delete_list_by_name(self, list_name):
        if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف القائمة '{list_name}'؟"):
            del self.lists_data[list_name]
            save_lists(self.lists_data)
            list_folder = os.path.join(MAIN_FOLDER, list_name)
            if os.path.exists(list_folder):
                shutil.rmtree(list_folder)
            if list_name in self.lists_order:
                self.lists_order.remove(list_name)
            if list_name in self.lists_colors:
                del self.lists_colors[list_name]
            messagebox.showinfo("نجاح", f"تم حذف القائمة '{list_name}' بنجاح!")
            self.save_config()
            self.refresh_lists()

    def export_list_by_name(self, list_name):
        export_data = self.lists_data[list_name]
        export_filename = f"shared_{list_name}.json"
        with open(export_filename, "w", encoding="utf-8") as f:
            json.dump({list_name: export_data}, f, ensure_ascii=False, indent=4)
        messagebox.showinfo("نجاح", f"تم تصدير القائمة إلى الملف: {export_filename}")

    def change_list_color_by_name(self, list_name):
        # التأكد من أن اسم القائمة موجود ضمن بيانات القوائم
        if list_name not in self.lists_data:
            messagebox.showerror("خطأ", "القائمة غير موجودة.")
            return
        # فتح مربع اختيار اللون
        color = colorchooser.askcolor(title="اختر لون القائمة")[1]
        if color:
            self.lists_colors[list_name] = color
            self.save_config()
            self.refresh_lists()


    def open_progress_by_name(self, list_name):
        tasks = self.lists_data[list_name]
        ProgressWindow(self, list_name, tasks)


    def refresh_lists(self):
        query = self.search_entry.get().lower()
        self.lists_listbox.delete(0, tk.END)
        for list_name in self.lists_order:
            if list_name in self.lists_data and (query in list_name.lower() or any(query in task_obj["task"].lower() for task_obj in self.lists_data[list_name])):
                self.lists_listbox.insert(tk.END, list_name)
                index = self.lists_listbox.size() - 1
                if list_name in self.lists_colors:
                    self.lists_listbox.itemconfig(index, bg=self.lists_colors[list_name])
                else:
                    default_bg = self.bg_value if self.bg_type=="color" and self.bg_value else ("black" if self.dark_mode else "white")
                    self.lists_listbox.itemconfig(index, bg=default_bg)

    def open_list(self):
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة.")
            return
        list_name = self.lists_listbox.get(selection[0])
        tasks = self.lists_data[list_name]
        DailyTrackerApp(self, list_name, tasks)

    def create_new_list(self):
        CreateListWindow(self)

    def edit_list(self):
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة لتعديلها.")
            return
        list_name = self.lists_listbox.get(selection[0])
        tasks = self.lists_data[list_name]
        CreateListWindow(self, list_name, tasks)

    def delete_list(self):
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة لحذفها.")
            return
        list_name = self.lists_listbox.get(selection[0])
        if messagebox.askyesno("تأكيد", f"هل أنت متأكد من حذف القائمة '{list_name}'؟"):
            del self.lists_data[list_name]
            save_lists(self.lists_data)
            list_folder = os.path.join(MAIN_FOLDER, list_name)
            if os.path.exists(list_folder):
                shutil.rmtree(list_folder)
            if list_name in self.lists_order:
                self.lists_order.remove(list_name)
            if list_name in self.lists_colors:
                del self.lists_colors[list_name]
            messagebox.showinfo("نجاح", f"تم حذف القائمة '{list_name}' بنجاح!")
            self.save_config()
            self.refresh_lists()

    def export_list(self):
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة للتصدير.")
            return
        list_name = self.lists_listbox.get(selection[0])
        export_data = self.lists_data[list_name]
        export_filename = f"shared_{list_name}.json"
        with open(export_filename, "w", encoding="utf-8") as f:
            json.dump({list_name: export_data}, f, ensure_ascii=False, indent=4)
        messagebox.showinfo("نجاح", f"تم تصدير القائمة إلى الملف: {export_filename}")

    def open_progress(self):
        selection = self.lists_listbox.curselection()
        if not selection:
            messagebox.showerror("خطأ", "يرجى اختيار قائمة لعرض التقدم.")
            return
        list_name = self.lists_listbox.get(selection[0])
        tasks = self.lists_data[list_name]
        ProgressWindow(self, list_name, tasks)

# ---------------------------
# نافذة إنشاء/تعديل القائمة مع دعم إعادة ترتيب المهام
# ---------------------------
class CreateListWindow(tk.Toplevel):
    def __init__(self, master, list_name=None, tasks=None):
        super().__init__(master)
        self.title("إنشاء / تعديل قائمة")
        self.geometry("450x500")
        self.master = master
        self.original_list_name = list_name
        self.tasks = tasks if tasks else []
        self.task_entries = []
        self.drag_data = {"widget": None, "y": 0}
        self.create_widgets()

    def create_widgets(self):
        motivational_messages = [
            "بسم الله، انطلاقة قوية! 🚀",
            "اللهم بارك لي في وقتي ✨",
            "سعيٌ وعمل.. توفيقٌ وأمل 💪",
            "ذكرٌ لله.. وبركة في الإنجاز 🌿",
            "باسمك اللهم نبدأ 💡",
            "يا رب سدد خُطاي 🤲",
            "العمل عبادة.. بسم الله!",
            "اللهم اجعل هذا اليوم خيرًا وبركة 🌞",
            "مع كل إنجاز.. الحمد لله ❤️",
            "لا كسل مع ذكر الله! 🚀"
        ]
        motivation = random.choice(motivational_messages)
        tk.Label(self, text=motivation, font=("Arial", 12, "italic"), fg="green").pack(pady=5)

        tk.Label(self, text="اسم القائمة:", font=("Arial", 12)).pack(pady=5)
        self.list_name_entry = tk.Entry(self, font=("Arial", 12))
        self.list_name_entry.pack(pady=5, fill="x", padx=20)
        if self.original_list_name:
            self.list_name_entry.insert(0, self.original_list_name)

        tk.Label(self, text="المهام:", font=("Arial", 12)).pack(pady=5)
        self.tasks_frame = tk.Frame(self)
        self.tasks_frame.pack(pady=5, fill="both", expand=True)
        tk.Button(self, text="إضافة مهمة", command=self.add_task_entry).pack(pady=10)
        for task_obj in self.tasks:
            self.add_task_entry(task_obj.get("task", ""), task_obj.get("priority", "متوسطة"))
        tk.Button(self, text="حفظ", command=self.save_list).pack(pady=10)
    
        tk.Label(self, text="created by: meedoasadel@gmail.com", font=("Arial", 20, "bold"), fg="blue").pack(side="bottom", pady=5)

    def add_task_entry(self, task_text="", priority="متوسطة"):
        frame = tk.Frame(self.tasks_frame, bd=1, relief="groove")
        frame.pack(fill="x", padx=10, pady=5)
        entry = tk.Entry(frame, font=("Arial", 12))
        entry.pack(side="left", fill="x", expand=True)
        entry.insert(0, task_text)
        priority_var = tk.StringVar(value=priority)
        option_menu = tk.OptionMenu(frame, priority_var, "عالية", "متوسطة", "منخفضة")
        option_menu.pack(side="left", padx=5)
        remove_btn = tk.Button(frame, text="حذف", command=lambda: self.remove_task_entry(frame))
        remove_btn.pack(side="left", padx=5)
        frame.bind("<Button-1>", self.on_drag_start)
        frame.bind("<B1-Motion>", self.on_drag_motion)
        frame.bind("<ButtonRelease-1>", self.on_drag_stop)
        self.task_entries.append({"frame": frame, "entry": entry, "priority": priority_var})

    def on_drag_start(self, event):
        widget = event.widget
        self.drag_data["widget"] = widget
        self.drag_data["y"] = event.y_root

    def on_drag_motion(self, event):
        dy = event.y_root - self.drag_data["y"]
        widget = self.drag_data["widget"]
        index = self.tasks_frame.winfo_children().index(widget)
        new_index = index
        if dy < -30 and index > 0:
            new_index = index - 1
        elif dy > 30 and index < len(self.tasks_frame.winfo_children()) - 1:
            new_index = index + 1
        if new_index != index:
            widget.pack_forget()
            widget.pack(in_=self.tasks_frame, before=self.tasks_frame.winfo_children()[new_index], fill="x", padx=10, pady=5)
            self.drag_data["y"] = event.y_root
    
    def on_drag_stop(self, event):
        self.drag_data = {"widget": None, "y": 0}

    def remove_task_entry(self, frame):
        for item in self.task_entries:
            if item["frame"] == frame:
                self.task_entries.remove(item)
                break
        frame.destroy()

    def save_list(self):
        list_name = self.list_name_entry.get().strip()
        if not list_name:
            messagebox.showerror("خطأ", "الرجاء إدخال اسم القائمة.")
            return

        tasks = []
        for item in self.task_entries:
            task_text = item["entry"].get().strip()
            if task_text:
                tasks.append({"task": task_text, "priority": item["priority"].get()})
        if not tasks:
            messagebox.showerror("خطأ", "يجب إضافة مهام على الأقل.")
            return

        if self.original_list_name and self.original_list_name != list_name:
            if list_name in self.master.lists_data:
                messagebox.showerror("خطأ", "هذه القائمة موجودة بالفعل.")
                return
            del self.master.lists_data[self.original_list_name]
            if self.original_list_name in self.master.lists_order:
                index = self.master.lists_order.index(self.original_list_name)
                self.master.lists_order[index] = list_name
            if self.original_list_name in self.master.lists_colors:
                self.master.lists_colors[list_name] = self.master.lists_colors.pop(self.original_list_name)
        elif not self.original_list_name and list_name in self.master.lists_data:
            messagebox.showerror("خطأ", "هذه القائمة موجودة بالفعل.")
            return

        self.master.lists_data[list_name] = tasks
        save_lists(self.master.lists_data)
        if list_name not in self.master.lists_order:
            self.master.lists_order.append(list_name)
        messagebox.showinfo("نجاح", "تم حفظ القائمة بنجاح!")
        self.master.save_config()
        self.master.refresh_lists()
        self.destroy()

# ---------------------------
# نافذة تتبع الأداء اليومي للقائمة
# ---------------------------
class DailyTrackerApp(tk.Toplevel):
    def __init__(self, master, list_name, tasks):
        super().__init__(master)
        self.list_name = list_name
        self.tasks = tasks
        self.list_folder = os.path.join(MAIN_FOLDER, list_name)
        self.excel_file = os.path.join(self.list_folder, "daily_progress.xlsx")
        self.title(f"تتبع الأداء - {list_name}")
        self.geometry("700x650")
        self.create_widgets()

    def create_widgets(self):
        today_str = datetime.date.today().strftime("%Y-%m-%d")
        tk.Label(self, text=f"تسجيل البيانات ليوم: {today_str}", font=("Arial", 16)).pack(pady=10)
        
        motivational_messages = [
            "بسم الله، انطلاقة قوية! 🚀",
            "اللهم بارك لي في وقتي ✨",
            "سعيٌ وعمل.. توفيقٌ وأمل 💪",
            "ذكرٌ لله.. وبركة في الإنجاز 🌿",
            "باسمك اللهم نبدأ 💡",
            "يا رب سدد خُطاي 🤲",
            "العمل عبادة.. بسم الله!",
            "اللهم اجعل هذا اليوم خيرًا وبركة 🌞",
            "مع كل إنجاز.. الحمد لله ❤️",
            "لا كسل مع ذكر الله! 🚀"
        ]
        motivation = random.choice(motivational_messages)
        tk.Label(self, text=motivation, font=("Arial", 12, "italic"), fg="green").pack(pady=5)
        
        self.task_vars = {}
        self.comment_vars = {}
        tasks_frame = tk.Frame(self)
        tasks_frame.pack(pady=10, fill="both", expand=True)
        for task_obj in self.tasks:
            task = task_obj["task"]
            row = tk.Frame(tasks_frame)
            row.pack(fill="x", pady=5, padx=10)
            var = tk.BooleanVar()
            tk.Checkbutton(row, variable=var).pack(side="left", padx=5)
            tk.Label(row, text=task, font=("Arial", 12)).pack(side="left", padx=5)
            tk.Label(row, text=f"({task_obj.get('priority', 'متوسطة')})", font=("Arial", 10), fg="blue").pack(side="left", padx=5)
            self.task_vars[task] = var
            comment_var = tk.StringVar()
            tk.Entry(row, textvariable=comment_var, width=40, font=("Arial", 12)).pack(side="left", padx=5)
            self.comment_vars[task] = comment_var

        tk.Button(self, text="حفظ البيانات", command=self.save_data).pack(pady=10)
        rep_frame = tk.Frame(self)
        rep_frame.pack(pady=5)
        tk.Button(rep_frame, text="توليد التقرير الأسبوعي", command=self.generate_weekly_report).pack(side="left", padx=5)
        tk.Button(rep_frame, text="توليد التقرير الشهري", command=self.generate_monthly_report).pack(side="left", padx=5)
        tk.Button(rep_frame, text="توليد تقرير PDF", command=self.generate_pdf_report).pack(side="left", padx=5)
        
        tk.Label(self, text="created by: meedoasadel@gmail.com", font=("Arial", 20, "bold"), fg="blue").pack(side="bottom", pady=5)

    def save_data(self):
        today_str = datetime.date.today().strftime("%Y-%m-%d")
        data = {"التاريخ": today_str, "التعليقات": {}}
        for task_obj in self.tasks:
            task = task_obj["task"]
            data[task] = self.task_vars[task].get()
            data["التعليقات"][task] = self.comment_vars[task].get()
        if os.path.exists(self.excel_file):
            df = pd.read_excel(self.excel_file)
        else:
            df = pd.DataFrame()
        new_row = {"التاريخ": data["التاريخ"]}
        for task_obj in self.tasks:
            task = task_obj["task"]
            status = '✔' if data.get(task, False) else '✖️'
            new_row[task] = status
            new_row[f"{task}_تعليق"] = data["التعليقات"].get(task, "")
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        if not os.path.exists(self.list_folder):
            os.makedirs(self.list_folder)
    
        df.to_excel(self.excel_file, index=False)
        messagebox.showinfo("نجاح", "تم حفظ البيانات بنجاح!")

    def generate_weekly_report(self):
        self.generate_report(period="weekly")

    def generate_monthly_report(self):
        self.generate_report(period="monthly")

    def generate_report(self, period="weekly"):
        if not os.path.exists(self.excel_file):
            messagebox.showerror("خطأ", "لا توجد بيانات.")
            return
        df = pd.read_excel(self.excel_file)
        df['التاريخ'] = pd.to_datetime(df['التاريخ'])
        today = pd.Timestamp.today().normalize()
        if period == "weekly":
            start_period = today - pd.Timedelta(days=today.weekday())
            folder_name = "Weekly_Reports"
        else:
            start_period = today.replace(day=1)
            folder_name = "Monthly_Reports"
        period_data = df[df['التاريخ'] >= start_period]
        if period_data.empty:
            messagebox.showinfo("تقرير", f"لا توجد بيانات للفترة المحددة ({period}).")
            return
        summary = {}
        for task_obj in self.tasks:
            task = task_obj["task"]
            count = period_data[task].apply(lambda x: 1 if x == "✔" else 0).sum()
            comments_series = period_data[f"{task}_تعليق"].dropna().astype(str)
            comments = "; ".join([c for c in comments_series if c.strip() != ""])
            summary[task] = {"count": count, "comments": comments}
        document = Document()
        document.add_heading(f"تقرير {period} - {self.list_name}", 0)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "المهمة"
        hdr_cells[1].text = "عدد مرات القيام"
        hdr_cells[2].text = "التعليقات"
        for task_obj in self.tasks:
            task = task_obj["task"]
            row_cells = table.add_row().cells
            row_cells[0].text = task
            row_cells[1].text = str(summary[task]["count"])
            row_cells[2].text = summary[task]["comments"]
        report_date = today.strftime("%Y-%m-%d")
        report_folder = os.path.join(self.list_folder, folder_name)
        if not os.path.exists(report_folder):
            os.makedirs(report_folder)
        word_report_file = os.path.join(report_folder, f"{period}_report_{report_date}.docx")
        document.save(word_report_file)
        messagebox.showinfo("تقرير", f"تم حفظ تقرير {period} في:\n{word_report_file}")

    def generate_pdf_report(self):
        if not os.path.exists(self.excel_file):
            messagebox.showerror("خطأ", "لا توجد بيانات.")
            return
        df = pd.read_excel(self.excel_file)
        df['التاريخ'] = pd.to_datetime(df['التاريخ'])
        today = pd.Timestamp.today().normalize()
        start_date = today - pd.Timedelta(days=7)
        period_data = df[df['التاريخ'] >= start_date]
        if period_data.empty:
            messagebox.showinfo("تقرير", "لا توجد بيانات للفترة المحددة.")
            return
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', '', 14)
         
        title = reshape_arabic_text(f"تقرير PDF للمهام - {self.list_name}")
        pdf.cell(200, 10, txt=title, ln=True, align="C")
        pdf.ln(10)
         
        for task_obj in self.tasks:
            task_original = task_obj["task"]
            count = period_data[task_original].apply(lambda x: 1 if x == "✔" else 0).sum()
            comments_series = period_data[f"{task_original}_تعليق"].dropna().astype(str)
            comments = "; ".join([c for c in comments_series if c.strip() != ""])
    
            pdf.set_font('DejaVu', '', 12)
            pdf.cell(0, 10, txt=reshape_arabic_text(f"المهمة: {task_original}"), ln=True)
            pdf.cell(0, 10, txt=reshape_arabic_text(f"عدد المرات: {count}"), ln=True)
            pdf.cell(0, 10, txt=reshape_arabic_text(f"التعليقات: {comments}"), ln=True)
    
            pdf.ln(10)
 
        pdf_folder = os.path.join(self.list_folder, "PDF_Reports")
        if not os.path.exists(pdf_folder):
            os.makedirs(pdf_folder)
        pdf_file = os.path.join(pdf_folder, f"pdf_report_{today.strftime('%Y-%m-%d')}.pdf")
        pdf.output(pdf_file)
        messagebox.showinfo("تقرير", f"تم حفظ تقرير PDF في:\n{pdf_file}")

# ---------------------------
# بدء تشغيل التطبيق
# ---------------------------
if __name__ == "__main__":
    if not os.path.exists(MAIN_FOLDER):
        os.makedirs(MAIN_FOLDER)
    app = TaskManagerApp()
    app.mainloop()
